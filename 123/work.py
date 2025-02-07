
from openpyxl import Workbook
from openpyxl.utils.dataframe import dataframe_to_rows
from openpyxl.styles import Alignment
import pandas as pd
from docx import Document



filename = 'roob/Карты ОПР.docx'

# Открываем документ Word
document = Document(filename)

# Извлекаем таблицы из документа
tables = document.tables

# Создаем новый Excel Workbook
wb = Workbook()

# Флаг для первого листа
first_sheet = True

# Обрабатываем каждую таблицу и сохраняем ее данные в DataFrame
for idx, table in enumerate(tables):
    df = [['' for _ in range(len(table.columns))] for _ in range(len(table.rows))]
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            df[i][j] = cell.text.strip()
    df_table = pd.DataFrame(df)

    # Создаем новый лист для каждой таблицы
    if first_sheet:
        ws = wb.active
        ws.title = f'Table_{idx + 1}'
        first_sheet = False
    else:
        ws = wb.create_sheet(title=f'Table_{idx + 1}')

    # Записываем данные DataFrame в лист Excel
    for r in dataframe_to_rows(df_table, index=False, header=False):
        ws.append(r)

    # Применяем выравнивание для всех ячеек
    for row in ws.iter_rows():
        for cell in row:
            cell.alignment = Alignment(horizontal='center', vertical='center')

    # Пример объединения ячеек (если требуется)

    if idx % 3 == 0:
        ws.merge_cells('B1:H1')  # Пример объединения ячеек
        ws.merge_cells('I1:J1')
        ws.merge_cells('B2:D2')
        ws.merge_cells('F2:M2')
        ws.merge_cells('A3:C3')
        ws.merge_cells('D3:M3')
        ws.merge_cells('A4:C4')
        ws.merge_cells('D4:M4')
        ws.merge_cells('A5:F5')
        ws.merge_cells('G5:M5')
        ws.merge_cells('A6:F6')
        ws.merge_cells('H6:I6')
        ws.merge_cells('J6:K6')
        ws.merge_cells('L6:M6')
        ws.merge_cells('A7:F7')
        ws.merge_cells('A8:F8')
        ws.merge_cells('A9:F9')
        ws.merge_cells('A10:F10')
        ws.merge_cells('A11:B11')
        ws.merge_cells('C11:M11')
        ws.merge_cells('C12:M12')
        ws.merge_cells('C13:M13')
        ws.merge_cells('C14:M14')
        ws.merge_cells('C15:M15')
        ws.merge_cells('D3:T3')
    elif idx % 3 == 1:
        ws.merge_cells('A1:K1')
        ws.merge_cells('L1:R1')
        ws.merge_cells('A2:C2')
        ws.merge_cells('A3:C3')
        ws.merge_cells('A2:C3')
        ws.merge_cells('D2:D3')
        ws.merge_cells('E2:F3')
        ws.merge_cells('G2:G3')
        ws.merge_cells('H2:I3')
        ws.merge_cells('J2:K3')
        ws.merge_cells('l2:P2')
        ws.merge_cells('Q2:R3')
        ws.merge_cells('A4:P4')



        ws.merge_cells('A24:P24')
        ws.merge_cells('A34:P34')
        ws.merge_cells('B5:C5')
        ws.merge_cells('B7:C7')
        ws.merge_cells('B10:C10')
        ws.merge_cells('B11:C12')
        ws.merge_cells('B12:C12')
        ws.merge_cells('B13:C13')
        ws.merge_cells('B14:C14')
        ws.merge_cells('B15:C15')
        ws.merge_cells('B16:C17')
        ws.merge_cells('B17:C17')
        ws.merge_cells('B18:C18')
        ws.merge_cells('B19:C19')
        ws.merge_cells('B20:C20')
        ws.merge_cells('B22:C22')
        ws.merge_cells('B23:C23')
        ws.merge_cells('B25:C25')
        ws.merge_cells('B26:C26')
        ws.merge_cells('B27:C27')
        ws.merge_cells('B28:C28')
        ws.merge_cells('B29:C29')
        ws.merge_cells('B30:C30')
        ws.merge_cells('B31:C31')
        ws.merge_cells('B32:C32')
        ws.merge_cells('B33:C33')
        ws.merge_cells('B35:C35')
        ws.merge_cells('B36:C36')
        ws.merge_cells('B37:C37')
        ws.merge_cells('B38:C38')
        ws.merge_cells('B39:C39')
        ws.merge_cells('B40:C40')
        ws.merge_cells('F5:G5')
        ws.merge_cells('F6:G6')
        ws.merge_cells('F7:G7')
        ws.merge_cells('F8:G8')
        ws.merge_cells('F9:G9')
        ws.merge_cells('F10:G10')
        ws.merge_cells('F11:G11')
        ws.merge_cells('F12:G12')
        ws.merge_cells('F13:G13')
        ws.merge_cells('F14:G14')
        ws.merge_cells('F15:G15')
        ws.merge_cells('F16:G16')
        ws.merge_cells('F17:G17')
        ws.merge_cells('F18:G18')
        ws.merge_cells('F19:G19')
        ws.merge_cells('F20:G20')
        ws.merge_cells('F21:G21')
        ws.merge_cells('F22:G22')
        ws.merge_cells('F23:G23')
        ws.merge_cells('F25:G25')
        ws.merge_cells('F26:G26')
        ws.merge_cells('F27:G27')
        ws.merge_cells('F28:G28')
        ws.merge_cells('F29:G29')
        ws.merge_cells('F30:G30')
        ws.merge_cells('F31:G31')
        ws.merge_cells('F32:G32')
        ws.merge_cells('F33:G33')

    #elif idx % 3 == 2:


# Сохраняем файл Excel
output_filename = 'roob2/output.xlsx'
wb.save(output_filename)

print(f"Data has been saved to {output_filename}")

import pandas as pd
import re

# Path to the uploaded Excel file
file_path = 'roob2/output.xlsx'

# Reading all sheets into a dictionary
all_sheets = pd.read_excel(file_path, sheet_name=None)

# Function to extract company names and professions
def extract_data(text):
    # Регулярное выражение для поиска нужного формата
    pattern = re.compile(r'(работники|профессий|персонал|состав|сотрудники):\s*([^.:]*?)\.\s*(.*)')
    matches = pattern.findall(text)

    companies = []
    professions = []

    for match in matches:
        company = match[1].strip()  # Извлекаем название компании
        profession_text = match[2].strip()  # Извлекаем текст профессии

        # Если company не пустая, добавляем ее в список компаний
        if company:
            companies.append(company)

        # Если profession_text не пустой, добавляем текст профессии
        if profession_text:
            # Разделяем текст по точке и берем последний элемент
            profession = profession_text.strip()
            if profession:
                professions.append(profession)

    return companies, professions

# List to store all pairs of (company, profession)
company_profession_list = []
print(company_profession_list )
# Processing all sheets
for sheet_name, data in all_sheets.items():
    for cell in data.to_numpy().flatten():
        if isinstance(cell, str):
            companies, professions = extract_data(cell)
            if companies and professions:
                for company in companies:
                    for profession in professions:
                        clean_company = company.strip()
                        clean_profession = profession.strip()
                        if clean_profession and clean_profession.startswith(clean_company + ", "):
                            clean_profession = clean_profession.replace(clean_company + ", ", "").strip()
                        if clean_profession:
                            # Adding each pair (company, profession) to the list
                            company_profession_list.append({'Company': clean_company, 'Profession': clean_profession})

# Creating DataFrame from the list of dictionaries
df = pd.DataFrame(company_profession_list)

# Saving DataFrame to a new Excel file
output_file_path = 'roob2/company_professions.xlsx'
df.to_excel(output_file_path, index=False)

# Displaying the DataFrame to the user











# Функция для извлечения всех числовых значений из документа
def extract_numerical_values(doc_path):
    # Загрузка документа
    doc = Document(doc_path)

    # Список для хранения извлечённых номеров
    extracted_numbers = []

    # Проход по абзацам и поиск числовых значений
    for paragraph in doc.paragraphs:
        matches = re.findall(r'\s*№\s*(\d+)\.?\s*', paragraph.text)
        extracted_numbers.extend(matches)

    return extracted_numbers
print()
# Путь к документу
doc_path = "roob/Карты ОПР.docx"

# Извлечение числовых значений
numerical_values = extract_numerical_values(doc_path)
print(f'Извлеченные числовые значения: {numerical_values}')
count = len(numerical_values)
print(count)

from docx import Document
import copy


# Функция для копирования содержимого указанное количество раз
def copy_content_multiple_times(src_doc_path, dest_doc_path, output_doc_path, num_copies):
    # Загрузка исходного и конечного документов
    src_doc = Document(src_doc_path)
    dest_doc = Document(dest_doc_path)

    # Копирование содержимого указанное количество раз
    for _ in range(num_copies):
        for element in src_doc.element.body:
            dest_doc.element.body.append(copy.deepcopy(element))

    # Сохранение конечного документа
    dest_doc.save(output_doc_path)


# Пути к вашим документам
src_doc_path = "План_мероприятий_по_снижению_искл_проф_рисков_ООО_СТАРТСПЕЙС.docx"
dest_doc_path = "План_мероприятий_по_снижению_искл_проф_рисков_ООО_СТАРТСПЕЙС1.docx"
output_doc_path = "final-document.docx"

# Количество копий
num_copies = (count-1)

# Вызов функции копирования содержимого указанное количество раз
copy_content_multiple_times(src_doc_path, dest_doc_path, output_doc_path, num_copies)



from docx import Document


def replace_card_numbers(doc_path, output_path):
    # Открываем документ
    doc = Document(doc_path)

    # Инициализация счетчика карт
    card_counter = 1

    # Проходимся по всем абзацам в документе
    for paragraph in doc.paragraphs:
        if "Карта №1." in paragraph.text:
            # Замена текста на "Карта №<номер>"
            paragraph.text = paragraph.text.replace("Карта №1.", f"Карта №{card_counter}.")
            card_counter += 1

    # Сохраняем измененный документ
    doc.save(output_path)


# Путь к вашему документу
input_doc_path = "final-document.docx"
output_doc_path = "roob/План_мероприятий_по_снижению_искл_проф_рисков_ООО_СТАРТСПЕЙС.docx"

# Вызов функции для замены номеров карт
replace_card_numbers(input_doc_path, output_doc_path)


# Путь к вашему Excel файлу
excel_file_path = 'roob2/company_professions.xlsx'

# Чтение данных из Excel файла
df = pd.read_excel(excel_file_path)

# Добавление новой колонки с числовыми значениями из Word-документа
# Дополнение числовыми значениями, если их меньше чем строк в DataFrame
if len(numerical_values) < len(df):
    numerical_values.extend([None] * (len(df) - len(numerical_values)))

df['Numerical Values'] = numerical_values[:len(df)]

# Запись DataFrame в новый Excel файл
output_file_path = 'roob2/company_professions.xlsx'
df.to_excel(output_file_path, index=False)

print(f'Данные обновлены и сохранены в файл: {output_file_path}')


# Функция для сбора и обработки данных из ячеек F5:G32
def collect_and_process_data(sheet):
    data = []
    for row in range(1, 70):  # Диапазон от 5 до 32 включительно
        f_value = sheet[f'F{row}'].value
        g_value = sheet[f'G{row}'].value
        data.append((f_value, g_value))

    flattened_data = []
    for item in data:
        if item[0] is not None:
            cleaned_item = item[0].replace(',', '')
            flattened_data.extend(cleaned_item.split())
        if item[1] is not None:
            cleaned_item = item[1].replace(',', '')
            flattened_data.extend(cleaned_item.split())

    return flattened_data

# Функция для загрузки данных из документа Word


def load_data_from_docx(doc_path):
    doc = Document(doc_path)
    extracted_data = []

    for table in doc.tables:
        for row in table.rows:
            code = row.cells[1].text
            extracted_data.append(code)

    return extracted_data


# Функция для сохранения данных в новый Excel файл
def save_data_to_excel(extracted_data, processed_file_path, new_file_path):
    processed_workbook = openpyxl.load_workbook(processed_file_path)
    processed_sheet = processed_workbook.active

    new_workbook = openpyxl.Workbook()
    new_sheet = new_workbook.active
    new_sheet.title = "Matches"

    # Создаем множество для быстрого поиска
    extracted_data_set = set(extracted_data)

    for col_idx, col in enumerate(
            processed_sheet.iter_cols(min_col=1, max_col=processed_sheet.max_column, values_only=True), start=1):
        matches = []
        for cell_value in col:
            if cell_value:
                # Разделяем элементы в ячейке, если они содержат точку с запятой
                cell_values = cell_value.split(';')
                # Проверяем каждый элемент на наличие в извлеченных данных
                matches.extend([val for val in cell_values if val in extracted_data_set])

        # Записываем совпадения в новый файл
        for row_idx, match in enumerate(matches, start=1):
            new_sheet.cell(row=row_idx, column=col_idx, value=match)

    new_workbook.save(new_file_path)
    print(f"Совпадения сохранены в '{new_file_path}'.")

# Функция для обработки данных из исходного файла Excel и сохранения их в новый файл
def process_and_save_excel_data(file_path, output_path):
    workbook = openpyxl.load_workbook(file_path)
    sheet_names = workbook.sheetnames

    new_workbook = openpyxl.Workbook()
    new_sheet = new_workbook.active
    new_sheet.title = "Collected Data"

    column_index = 1
    for i in range(1, len(sheet_names), 3):
        sheet_name = sheet_names[i]
        sheet = workbook[sheet_name]
        result_array = collect_and_process_data(sheet)

        for row_index, value in enumerate(result_array, start=1):
            new_sheet.cell(row=row_index, column=column_index, value=value)

        column_index += 1

    new_workbook.save(output_path)
    print(f"Обработанные данные сохранены в '{output_path}'.")

file_path = 'roob2/output.xlsx'
processed_output_path = 'roob2/processed_output.xlsx'
process_and_save_excel_data(file_path, processed_output_path)

import openpyxl
from openpyxl.utils import range_boundaries

# Пути к файлам
input_file_path = 'roob2/processed_output.xlsx'
output_file_path = 'roob2/processed_output.xlsx'

try:
    # Загрузка книги Excel
    wb = openpyxl.load_workbook(input_file_path)
    sheet = wb.active

    # Получение максимального количества строк и столбцов
    max_row = sheet.max_row
    max_column = sheet.max_column

    # Создание списка для хранения слияний ячеек
    merged_cells = list(sheet.merged_cells.ranges)

    # Удаление текущих слияний ячеек
    for merged_cell in merged_cells:
        sheet.unmerge_cells(str(merged_cell))

    # Сдвиг всех строк на одну вниз
    for row in range(max_row, 0, -1):
        for col in range(1, max_column + 1):
            cell = sheet.cell(row=row, column=col)
            new_cell = sheet.cell(row=row+1, column=col)
            if not isinstance(cell, openpyxl.cell.cell.MergedCell):
                new_cell.value = cell.value
                cell.value = None

    # Восстановление слияний ячеек с новой позицией
    for merged_cell in merged_cells:
        min_col, min_row, max_col, max_row = range_boundaries(str(merged_cell))
        new_range = f"{openpyxl.utils.get_column_letter(min_col)}{min_row+1}:{openpyxl.utils.get_column_letter(max_col)}{max_row+1}"
        sheet.merge_cells(new_range)

    # Сохранение изменений в том же файле Excel
    wb.save(output_file_path)

    print("Файл успешно обработан и сохранен.")

except FileNotFoundError:
    print(f"Файл не найден по указанному пути: {input_file_path}")
except Exception as e:
    print(f"Произошла ошибка: {e}")


doc_path = 'roob/Сводный_реестр_опасностей_ООО_СТАРТСПЕЙС.docx'
extracted_data = load_data_from_docx(doc_path)
print("Извлеченные данные из документа Word:", extracted_data)

matched_output_path = 'roob2/matched_output.xlsx'
save_data_to_excel(extracted_data, processed_output_path, matched_output_path)


# Load the Excel file
file_path = 'roob2/matched_output.xlsx'  # Replace with your file path
df = pd.read_excel(file_path)

# Combine the data in each column into a single string, separated by a semicolon and space
combined_columns = df.apply(lambda col: '; '.join(col.dropna().astype(str)), axis=0)

# Convert the combined series to a DataFrame to arrange them vertically
combined_columns_vertical_df = pd.DataFrame(combined_columns, columns=['Combined'])

# Save the new DataFrame to a new Excel file
final_output_path = 'roob2/your_output_file.xlsx'  # Replace with your desired output file path
combined_columns_vertical_df.to_excel(final_output_path, index=False, header=False)

print(f"Combined vertical columns output saved to {final_output_path}")

# Открытие существующей книги Excel из указанной директории
input_file_path = 'roob2/matched_output.xlsx'
output_file_path = 'roob2/matched_output.xlsx'

wb = openpyxl.load_workbook(input_file_path)
sheet = wb.active

# Получение максимального количества строк и столбцов
max_row = sheet.max_row
max_column = sheet.max_column

# Сдвиг всех строк на одну вниз
for col in range(1, max_column + 1):
    for row in range(max_row, 0, -1):
        sheet.cell(row=row+1, column=col).value = sheet.cell(row=row, column=col).value
        sheet.cell(row=row, column=col).value = None

# Сохранение изменений в том же файле Excel
wb.save(output_file_path)



# Load the Excel file
file_path = 'roob2/matched_output.xlsx'  # Replace with your file path
df = pd.read_excel(file_path)

# Combine the data in each column into a single string, separated by a semicolon and space
combined_columns = df.apply(lambda col: '; '.join(col.dropna().astype(str)), axis=0)

# Convert the combined series to a DataFrame to arrange them vertically
combined_columns_vertical_df = pd.DataFrame(combined_columns, columns=['Combined'])

# Save the new DataFrame to a new Excel file
final_output_path = 'roob2/your_output_file.xlsx'  # Replace with your desired output file path
combined_columns_vertical_df.to_excel(final_output_path, index=False, header=False)

print(f"Combined vertical columns output saved to {final_output_path}")



# Load the Excel file
file_path = 'roob2/company_professions.xlsx'
xls = pd.ExcelFile(file_path)

# Load the data from the sheet
df = pd.read_excel(file_path, sheet_name='Sheet1')

# Add a new column 'Code' with placeholder values
if len(df.columns) < 5:
    df['Code'] = None
else:
    df.columns = list(df.columns[:4]) + ['Code']

# Save the modified DataFrame back to the same Excel file, overwriting the existing sheet
with pd.ExcelWriter(file_path, engine='openpyxl') as writer:
    df.to_excel(writer, index=False, sheet_name='Sheet1')

print("New column 'Code' added and saved to the same file.")



import pandas as pd

# Load the Excel files
file_path_1 = 'roob2/company_professions.xlsx'
file_path_2 = 'roob2/your_output_file.xlsx'

# Load data from both sheets
df1 = pd.read_excel(file_path_1, sheet_name='Sheet1')
df2 = pd.read_excel(file_path_2, sheet_name='Sheet1', header=None)

# Assuming df2 contains the data for the fifth column and has the same number of rows
df1['Code'] = df2.iloc[:, 0]

# Save the modified DataFrame back to the same Excel file, overwriting the existing sheet
with pd.ExcelWriter(file_path_1, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
    df1.to_excel(writer, index=False, sheet_name='Sheet1')

print("Data from the second file added to the fifth column and saved to the same file.")

import pandas as pd
import re

# Путь к исходному файлу Excel
file_path = 'roob2/output.xlsx'
xls = pd.ExcelFile(file_path)

# Список листов для обработки, начиная с третьего и далее каждый третий
sheet_names = xls.sheet_names[2::3]

# Извлечение и обработка данных
data_from_column_d = {}
for sheet in sheet_names:
    data = pd.read_excel(xls, sheet_name=sheet)
    # Предполагаем, что данные находятся в четвертом столбце, строка D3
    text = data.iloc[2, 3]  # Индексы для строки 3 и столбца D
    # Поиск чисел после слова "оценено"
    match = re.search(r"оценено (\d+)", text)
    if match:
        data_from_column_d[sheet] = int(match.group(1))
    else:
        data_from_column_d[sheet] = "Число не найдено"

# Создание DataFrame с результатами
df_numbers = pd.DataFrame(list(data_from_column_d.values()))

# Сохранение только колонки с числами, без заголовка
output_path = 'roob2/output_numbers.xlsx'
df_numbers.to_excel(output_path, index=False, header=False)

print(f"Data has been successfully saved to {output_path}")



# Пути к файлам Excel
file_path_1 = 'roob2/company_professions.xlsx'
file_path_2 = 'roob2/output_numbers.xlsx'

# Загрузка данных из обоих файлов
df1 = pd.read_excel(file_path_1, sheet_name='Sheet1')
df2 = pd.read_excel(file_path_2, header=None)  # Загрузка без заголовков

# Удостоверимся, что в df1 есть пять колонок
for i in range(len(df1.columns), 4):
    df1[f'Column {i + 1}'] = pd.NA

# Добавление данных из df2 в пятую колонку df1
df1['Column 5'] = df2[0]  # [0] указывает на первый столбец df2

# Сохранение изменённой таблицы df1 обратно в файл Excel, перезаписывая существующий лист
with pd.ExcelWriter(file_path_1, engine='openpyxl', mode='a', if_sheet_exists='replace') as writer:
    df1.to_excel(writer, index=False, sheet_name='Sheet1')

print("Data from the second file added to the fifth column and saved to the same file.")



import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Путь к вашему Excel файлу
excel_file_path = 'roob2/company_professions.xlsx'

# Чтение данных из Excel файла
df = pd.read_excel(excel_file_path)
excel_first_column = df.iloc[:, 0].dropna().astype(str).tolist()
excel_second_column = df.iloc[:, 1].dropna().astype(str).tolist()
# Допустим, что данные для третьего и четвертого столбца тоже находятся в таблице Excel
excel_third_column = df.iloc[:, 2].dropna().astype(str).tolist()
excel_phird_column = df.iloc[:, 3].dropna().astype(str).tolist()

# Путь к вашему Word файлу
word_file_path = 'roob/Перечень_вредных_и_опасных_производственных_факторов_ООО_СТАРТСПЕЙС.docx'

# Загружаем документ
doc = Document(word_file_path)

# Предполагаем, что таблица находится на первой странице и является второй таблицей в документе (индекс 1)
table = doc.tables[1]

# Функция для выравнивания текста в ячейке по центру
def center_text(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(10)  # Устанавливаем размер шрифта

# Проверка количества строк в таблице и в данных из Excel
max_rows = max(len(excel_first_column), len(excel_second_column), len(excel_third_column), len(excel_phird_column))
current_rows = len(table.rows) - 1  # исключая заголовок

if current_rows < max_rows:
    for _ in range(max_rows - current_rows):
        table.add_row()

# Проходим по каждой строке таблицы, начиная со второй (первая строка - заголовок)
for i, row in enumerate(table.rows[1:], start=0):  # начинаем с row[1:] для пропуска заголовка
    if i < len(excel_third_column):
        # Обновляем первую колонку (индекс 0)
        center_text(row.cells[0], excel_third_column[i])
    if i < len(excel_first_column):
        # Обновляем вторую колонку (индекс 1)
        center_text(row.cells[1], excel_first_column[i])
    if i < len(excel_second_column):
        # Обновляем третью колонку (индекс 2)
        center_text(row.cells[2], excel_second_column[i])
    if i < len(excel_phird_column):
        # Обновляем четвертую колонку (индекс 3)
        center_text(row.cells[3], excel_phird_column[i])

# Сохраняем изменения в новый документ
updated_word_file_path = 'gotfill/Перечень_вредных_и_опасных_производственных_факторов.docx'
doc.save(updated_word_file_path)

print(f'Данные обновлены и сохранены в файл: {updated_word_file_path}')



import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# Путь к вашему Excel файлу
excel_file_path = 'roob2/company_professions.xlsx'

# Чтение данных из Excel файла
df = pd.read_excel(excel_file_path)
excel_first_column = df.iloc[:, 0].dropna().astype(str).tolist()
excel_second_column = df.iloc[:, 1].dropna().astype(str).tolist()
# Допустим, что данные для третьего и четвертого столбца тоже находятся в таблице Excel
excel_third_column = df.iloc[:, 2].dropna().astype(str).tolist()
excel_forth_column = df.iloc[:, 4].dropna().astype(str).tolist()

# Путь к вашему Word файлу
word_file_path = 'roob/Сводная_ведомость_ООО_СТАРТСПЕЙС.docx'

# Загружаем документ
doc = Document(word_file_path)

# Предполагаем, что таблица находится на первой странице и является первой таблицей в документе
table = doc.tables[1]

# Функция для выравнивания текста в ячейке по центру
def center_text(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.size = Pt(10)  # Устанавливаем размер шрифта

# Проверка количества строк в таблице и в данных из Excel
max_rows = max(len(excel_first_column), len(excel_second_column), len(excel_third_column), len(excel_forth_column))
current_rows = len(table.rows) - 1  # исключая заголовок

if current_rows < max_rows:
    for _ in range(max_rows - current_rows):
        table.add_row()

# Проходим по каждой строке таблицы, начиная со второй (первая строка - заголовок)
for i, row in enumerate(table.rows[1:], start=0):  # начинаем с row[1:] для пропуска заголовка
    if i < len(excel_third_column):
        # Обновляем первую колонку (индекс 0)
        center_text(row.cells[0], excel_third_column[i])
    if i < len(excel_first_column):
        # Обновляем вторую колонку (индекс 1)
        center_text(row.cells[1], excel_first_column[i])
    if i < len(excel_second_column):
        # Обновляем третью колонку (индекс 2)
        center_text(row.cells[2], excel_second_column[i])
    if i < len(excel_forth_column):
        # Обновляем четвертую колонку (индекс 3)
        center_text(row.cells[3], excel_forth_column[i])

# Сохраняем изменения в новый документ
updated_word_file_path = 'gotfill/Сводная_ведомость.docx'
doc.save(updated_word_file_path)

print(f'Данные обновлены и сохранены в файл: {updated_word_file_path}')




import pandas as pd

# Загрузка файла Excel
file_path = 'roob2/output.xlsx'
excel_data = pd.ExcelFile(file_path)

# Получение всех имен листов
sheet_names = excel_data.sheet_names

# Создание объекта ExcelWriter для записи в новый файл Excel
output_file_path = 'roob2/combined_output.xlsx'
with pd.ExcelWriter(output_file_path, engine='xlsxwriter') as writer:
    # Итерация по каждой четвертой таблице
    for i in range(1, len(sheet_names), 3):
        sheet_name = sheet_names[i]
        table = excel_data.parse(sheet_name)

        # Фильтрация строк, где в колонке 'Unnamed: 4' содержится слово 'да'
        filtered_table = table[table['Unnamed: 4'].str.lower() == 'да']

        # Извлечение 6-й и 10-й колонок
        if not filtered_table.empty:
            extracted_data = filtered_table.iloc[:, [5, 9]]

            # Сохранение исходного порядка строк
            extracted_data = extracted_data.reset_index(drop=False)

            # Объединение строк с сохранением их уникальности
            combined_data = extracted_data.groupby(['Unnamed: 5', 'Unnamed: 9']).agg({
                'index': 'first'  # Сохранение исходного индекса для сортировки
            }).reset_index()

            # Сортировка по сохраненному исходному индексу
            combined_data = combined_data.sort_values('index').drop(columns='index')

            # Запись объединённых данных в новый лист в файле Excel
            combined_data.to_excel(writer, sheet_name=f'Combined_{i}', index=False)

print(f'Объединенные данные сохранены в файл: {output_file_path}')



import re
from docx import Document
import pandas as pd

# Load the Excel file
file_path = 'roob2/company_professions.xlsx'
df = pd.read_excel(file_path, sheet_name='Sheet1')

# Extracting the necessary columns
numerical_values = df['Numerical Values'].tolist()
company_names = df['Company'].tolist()
professions = df['Profession'].tolist()

# Function to replace subdivisions in the document text
def replace_subdivision(paragraph, company_names, index):
    subdivision_pattern = re.compile(r'(Подразделение:\s*)(.*)')
    return subdivision_pattern.sub(f'\\1{company_names[index]}', paragraph)

# Function to replace professions in the document text
def replace_profession(paragraph, professions, index):
    position_pattern = re.compile(r'(Должность:\s*)(.*)')
    return position_pattern.sub(f'\\1{professions[index]}', paragraph)

# Load the document
doc = Document('roob/План_мероприятий_по_снижению_искл_проф_рисков_ООО_СТАРТСПЕЙС.docx')

# Iterate through the paragraphs to replace card numbers, "Подразделение:", and "Должность:"
card_number_pattern = re.compile(r'Карта №\d+')

for i, paragraph in enumerate(doc.paragraphs):
    if card_number_pattern.search(paragraph.text):
        card_num_index = int(card_number_pattern.search(paragraph.text).group().split('№')[-1]) - 1
        paragraph.text = card_number_pattern.sub(f'Карта №{numerical_values[card_num_index]}', paragraph.text)
    if 'Подразделение:' in paragraph.text:
        paragraph.text = replace_subdivision(paragraph.text, company_names, card_num_index)
    if 'Должность:' in paragraph.text:
        paragraph.text = replace_profession(paragraph.text, professions, card_num_index)

# Save the document with the applied changes
doc.save('roob2/План_мероприятий_по_снижению_искл_проф_рисков_ООО_СТАРТСПЕЙС_ОБНОВЛЕННЫЙ.docx')



import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

excel_file_path = 'roob2/combined_output.xlsx'
word_file_path = 'roob2/План_мероприятий_по_снижению_искл_проф_рисков_ООО_СТАРТСПЕЙС_ОБНОВЛЕННЫЙ.docx'
updated_word_file_path = 'gotfill/План мероприятий по исключению или снижению уровней профессиональных рисков.docx'

# Чтение данных из Excel файла, включая все листы
excel_data = pd.read_excel(excel_file_path, sheet_name=None)

# Загружаем документ Word
doc = Document(word_file_path)

# Функция для выравнивания текста в ячейке по центру
def center_text(cell, text):
    cell.text = ''
    p = cell.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        run.font.size = Pt(10)  # Устанавливаем размер шрифта

# Функция для очистки текста в ячейке
def clear_text(cell):
    cell.text = ''

# Функция для удаления пустых строк
def remove_empty_rows(table):
    rows_to_delete = []
    for row in table.rows:
        is_empty = True
        for cell in row.cells:
            if cell.text.strip():
                is_empty = False
                break
        if is_empty:
            rows_to_delete.append(row)

    # Удаляем строки
    for row in rows_to_delete:
        row._element.getparent().remove(row._element)

# Функция для добавления недостающих колонок
def add_missing_columns(table, required_columns):
    current_columns = len(table.columns)
    if current_columns < required_columns:
        for row in table.rows:
            for _ in range(required_columns - current_columns):
                row.add_cell()

# Функция для добавления строк
def add_missing_rows(table, required_rows):
    current_rows = len(table.rows)
    for _ in range(required_rows - current_rows):
        table.add_row()

# Функция для добавления текста с маркерами
def add_bullet_points(cell, text_list):
    cell.text = ''
    for text in text_list:
        p = cell.add_paragraph()
        p.add_run('• ').bold = True
        p.add_run(text)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.size = Pt(10)  # Устанавливаем размер шрифта

# Перебираем все листы и таблицы, начиная со второй таблицы (индекс 1)
table_index = 1
for sheet_name, df in excel_data.items():
    # Получаем колонки как списки
    excel_first_column = df.iloc[:, 0].dropna().astype(str).tolist()
    excel_second_column = df.iloc[:, 1].dropna().astype(str).tolist()

    if table_index < len(doc.tables):
        table = doc.tables[table_index]
        table_index += 1

        # Удаляем пустые строки
        remove_empty_rows(table)

        # Добавляем недостающие колонки, если необходимо (не менее двух колонок)
        add_missing_columns(table, 2)

        # Добавляем недостающие строки, если необходимо
        add_missing_rows(table, len(excel_first_column) + 1)  # +1 для учета заголовка

        # Очищаем текущие данные в таблице
        for row in table.rows[1:]:  # начинаем с row[1:] для пропуска заголовка
            for cell in row.cells:
                clear_text(cell)

        # Заполняем таблицу новыми данными
        for i, row in enumerate(table.rows[1:], start=0):  # начинаем с row[1:] для пропуска заголовка
            if i < len(excel_first_column):
                # Обновляем первую колонку (индекс 0)
                center_text(row.cells[0], excel_first_column[i])
            if i < len(excel_second_column):
                # Обновляем вторую колонку (индекс 1) с маркерами
                bullet_points = excel_second_column[i].split('\n')
                add_bullet_points(row.cells[1], bullet_points)

# Сохраняем изменения в новый документ
doc.save(updated_word_file_path)
print(f'Dанные обновлены и сохранены в файл: {updated_word_file_path}')



import pandas as pd
from docx import Document

# Чтение документа
doc_path = "roob/Сводный_реестр_опасностей_ООО_СТАРТСПЕЙС.docx"
document = Document(doc_path)

# Извлечение данных из таблиц
data = {
    "Код": [],
    "Наименование опасности (вредный или опасный фактор)": []
}

# Парсинг таблиц из документа
for table in document.tables:
    for row in table.rows:
        cells = row.cells
        if len(cells) == 2:
            code = cells[1].text.strip()
            description = cells[0].text.strip()
            data["Код"].append(code)
            data["Наименование опасности (вредный или опасный фактор)"].append(description)

# Создание DataFrame
df = pd.DataFrame(data)

# Сохранение DataFrame в Excel
file_path = "roob2/integer.xlsx"
df.to_excel(file_path, index=False)

print(f"Данные успешно сохранены в {file_path}")

import pandas as pd

# Load the uploaded files
file1_path = 'roob2/integer.xlsx'
file2_path = 'roob2/matched_output.xlsx'

# Reading the first sheet of the first file
df1 = pd.read_excel(file1_path, sheet_name=0)

# Reading the first sheet of the second file
df2 = pd.read_excel(file2_path, sheet_name=0)

# Extracting the codes from the first dataframe
codes_df1 = df1.iloc[1:, 0].unique()

# Flattening the codes from the second dataframe
codes_df2 = df2.values.flatten()

# Finding common codes
common_codes = set(codes_df1) & set(codes_df2)

# Filtering the first dataframe to keep only the rows with common codes
filtered_df1 = df1[df1.iloc[:, 0].isin(common_codes)]

# Keeping only the columns with the codes and the names of the dangers
filtered_df1_final = filtered_df1[['Код', 'Наименование опасности (вредный или опасный фактор)']]

# Removing duplicate rows based on the 'Код' column
filtered_df1_final_no_duplicates = filtered_df1_final.drop_duplicates(subset='Код')

# Saving the filtered dataframe without duplicates to an Excel file
output_file_path = 'roob2/upgrade.xlsx'
filtered_df1_final_no_duplicates.to_excel(output_file_path, index=False)

output_file_path


import pandas as pd
from docx import Document
from docx.shared import Pt

# Путь к вашему Excel файлу
excel_file_path = 'roob2/upgrade.xlsx'

# Чтение данных из Excel файла
df = pd.read_excel(excel_file_path)
excel_first_column = df.iloc[:, 0].dropna().astype(str).tolist()
excel_second_column = df.iloc[:, 1].dropna().astype(str).tolist()

# Путь к вашему Word файлу
word_file_path = 'roob/Сводный_реестр_опасностей_ООО_СТАРТСПЕЙС.docx'

# Загружаем документ
doc = Document(word_file_path)

# Предполагаем, что таблица находится на первой странице и является третьей таблицей в документе (индекс 2)
table = doc.tables[2]

# Функция для обновления текста в ячейке
def update_text(cell, text):
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(10)  # Устанавливаем размер шрифта

# Проверка количества строк в таблице и в данных из Excel
max_rows = max(len(excel_first_column), len(excel_second_column))
current_rows = len(table.rows) - 1  # исключая заголовок

if current_rows < max_rows:
    for _ in range(max_rows - current_rows):
        table.add_row()

# Проходим по каждой строке таблицы, начиная со второй (первая строка - заголовок)
for i, row in enumerate(table.rows[1:], start=0):  # начинаем с row[1:] для пропуска заголовка
    if i < len(excel_first_column):
        # Обновляем первую колонку (индекс 0)
        update_text(row.cells[1], excel_first_column[i])
    if i < len(excel_second_column):
        # Обновляем вторую колонку (индекс 1)
        update_text(row.cells[0], excel_second_column[i])

# Сохраняем изменения в новый документ
updated_word_file_path = 'gotfill/Сводный_реестр_опасностей.docx'
doc.save(updated_word_file_path)

print(f'Данные обновлены и сохранены в файл: {updated_word_file_path}')

